import os
import logging
import re
import zipfile
import xml.etree.ElementTree as ET
from typing import List, Dict, Any
from datetime import datetime

from PyPDF2 import PdfReader
import cv2
import numpy as np
from pdf2image import convert_from_path
import pytesseract
from config import Config

logger = logging.getLogger(__name__)

# Check for optional dependencies
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning("python-docx not available. DOCX processing disabled.")

try:
    import markdown
    from bs4 import BeautifulSoup
    MARKDOWN_AVAILABLE = True
except ImportError:
    MARKDOWN_AVAILABLE = False
    logger.warning("markdown/beautifulsoup4 not available. Markdown processing will use plain text fallback.")

try:
    from striprtf.striprtf import rtf_to_text
    RTF_AVAILABLE = True
except ImportError:
    RTF_AVAILABLE = False
    logger.warning("striprtf not available. RTF processing disabled.")

try:
    import pdf2image
    import pytesseract
    from PIL import Image
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False
    logger.warning("OCR dependencies (pdf2image, pytesseract, pillow) not available. OCR fallback disabled for image-based PDFs.")

class DocumentProcessor:
    """Document processor for extracting text content from various file formats."""
    
    def __init__(self):
        self.config = Config()
        # Build supported extensions based on available dependencies
        base_extensions = {'.txt', '.md'}  # Always supported
        
        base_extensions.add('.pdf')  # PDF is always "supported" but may fall back if deps missing
        if DOCX_AVAILABLE:
            base_extensions.update({'.docx', '.doc'})
        if RTF_AVAILABLE:
            base_extensions.add('.rtf')
        
        # ODT is supported through manual parsing (no external dependency needed)
        base_extensions.add('.odt')
        
        self.supported_extensions = base_extensions
        logger.info(f"DocumentProcessor initialized with support for: {sorted(self.supported_extensions)}")
        if not OCR_AVAILABLE:
            logger.warning("OCR not available; image-based PDFs may fail to extract text.")
    
    def process_file(self, file_path: str, original_filename: str = None) -> str:
        """
        Process a file and extract its content.
        
        Args:
            file_path: Path to the temporary file
            original_filename: Original filename to determine file type
        
        Returns:
            Extracted text content
        """
        try:
            # Use original filename for type detection if provided
            filename_to_check = original_filename if original_filename else file_path
            file_extension = os.path.splitext(filename_to_check)[1].lower()
            
            logger.info(f"📄 Processing file type: {file_extension} (original: {original_filename})")
            
            if file_extension == '.pdf':
                return self._process_pdf(file_path)
            elif file_extension in ['.docx', '.doc']:
                if not DOCX_AVAILABLE:
                    raise ValueError("DOCX processing not available. Install python-docx: pip install python-docx")
                return self._process_docx(file_path)
            elif file_extension == '.txt':
                return self._process_txt(file_path)
            elif file_extension == '.md':
                return self._process_markdown(file_path)
            elif file_extension == '.rtf':
                if not RTF_AVAILABLE:
                    raise ValueError("RTF processing not available. Install striprtf: pip install striprtf")
                return self._process_rtf(file_path)
            elif file_extension == '.odt':
                return self._process_odt(file_path)
            else:
                logger.error(f"Unsupported file type: {file_extension}")
                raise ValueError(f"Unsupported file type: {file_extension}. Supported: {sorted(self.supported_extensions)}")
                
        except Exception as e:
            logger.error(f"Error processing file {file_path}: {str(e)}", exc_info=True)
            raise
    
    def _process_pdf(self, file_path: str) -> str:
        """Extract text from PDF files with configurable OCR fallback."""
        try:
            text_content = ""
            with open(file_path, 'rb') as file:
                pdf_reader = PdfReader(file)
                
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text:
                            text_content += f"\n--- Page {page_num + 1} ---\n"
                            text_content += page_text
                    except Exception as e:
                        logger.warning(f"Error extracting page {page_num + 1}: {e}")
                        continue
            
            # Check if OCR is needed based on configuration threshold
            min_threshold = getattr(self.config, 'OCR_MIN_TEXT_THRESHOLD', 100)
            enable_ocr = getattr(self.config, 'ENABLE_OCR', True)
            
            if enable_ocr and OCR_AVAILABLE and (not text_content.strip() or len(text_content.strip()) < min_threshold):
                logger.warning(
                    f"Minimal text extracted ({len(text_content.strip())} chars, "
                    f"threshold: {min_threshold}), attempting OCR..."
                )
                try:
                    ocr_text = self._extract_text_with_ocr(file_path)
                    if ocr_text and len(ocr_text.strip()) > len(text_content.strip()):
                        logger.info(f"✅ OCR extraction successful: {len(ocr_text)} characters")
                        text_content = ocr_text
                    else:
                        logger.warning("OCR did not improve text extraction, using original")
                except Exception as ocr_error:
                    logger.error(f"❌ OCR extraction failed: {str(ocr_error)}")
                    if not text_content.strip():
                        return "No text content could be extracted from this PDF file. OCR processing failed."
            elif enable_ocr and not OCR_AVAILABLE:
                logger.warning("OCR enabled but dependencies not available; skipping OCR fallback.")
            
            if not text_content.strip():
                return "No text content could be extracted from this PDF file."
            
            logger.info(f"✅ Extracted {len(text_content)} characters from PDF")
            return self._clean_text(text_content)
            
        except Exception as e:
            logger.error(f"Error processing PDF: {str(e)}", exc_info=True)
            raise

    
    def _process_docx(self, file_path: str) -> str:
        """Extract text from DOCX files."""
        try:
            doc = Document(file_path)
            text_content = ""
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_content += " | ".join(row_text) + "\n"
            
            if not text_content.strip():
                logger.warning("No text extracted from DOCX file")
                return "No text content could be extracted from this DOCX file."
            
            logger.info(f"✅ Extracted {len(text_content)} characters from DOCX")
            return self._clean_text(text_content)
            
        except Exception as e:
            logger.error(f"Error processing DOCX: {str(e)}", exc_info=True)
            raise
    
    def _process_txt(self, file_path: str) -> str:
        """Extract text from TXT files."""
        try:
            # Try different encodings
            encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        content = file.read()
                        
                    logger.info(f"✅ Read TXT file with {encoding} encoding: {len(content)} characters")
                    return self._clean_text(content)
                    
                except UnicodeDecodeError:
                    continue
            
            # If all encodings fail, try binary mode and decode with errors='ignore'
            with open(file_path, 'rb') as file:
                content = file.read().decode('utf-8', errors='ignore')
                
            logger.info(f"✅ Read TXT file with error handling: {len(content)} characters")
            return self._clean_text(content)
            
        except Exception as e:
            logger.error(f"Error processing TXT file: {str(e)}", exc_info=True)
            raise
    
    def _process_markdown(self, file_path: str) -> str:
        """Extract text from Markdown files."""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                md_content = file.read()
            
            if MARKDOWN_AVAILABLE:
                # Convert markdown to HTML then extract text
                html = markdown.markdown(md_content)
                soup = BeautifulSoup(html, 'html.parser')
                text_content = soup.get_text()
                logger.info(f"✅ Extracted {len(text_content)} characters from Markdown (with formatting)")
            else:
                # Fallback: treat as plain text but remove common markdown syntax
                text_content = self._clean_markdown_syntax(md_content)
                logger.info(f"✅ Extracted {len(text_content)} characters from Markdown (plain text fallback)")
            
            return self._clean_text(text_content)
            
        except Exception as e:
            logger.error(f"Error processing Markdown: {str(e)}", exc_info=True)
            # Fallback: treat as plain text
            return self._process_txt(file_path)
    
    def _process_rtf(self, file_path: str) -> str:
        """Extract text from RTF files."""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                rtf_content = file.read()
            
            text_content = rtf_to_text(rtf_content)
            
            logger.info(f"✅ Extracted {len(text_content)} characters from RTF")
            return self._clean_text(text_content)
            
        except Exception as e:
            logger.error(f"Error processing RTF: {str(e)}", exc_info=True)
            raise
    
    def _process_odt(self, file_path: str) -> str:
        """Extract text from ODT files using manual XML parsing."""
        try:
            text_content = ""
            
            with zipfile.ZipFile(file_path, 'r') as odt_file:
                # ODT files contain content.xml with the actual text
                try:
                    content_xml = odt_file.read('content.xml')
                    
                    # Parse XML and extract text
                    root = ET.fromstring(content_xml)
                    
                    # Remove namespace prefixes for easier parsing
                    for elem in root.iter():
                        if elem.tag.startswith('{'):
                            elem.tag = elem.tag.split('}', 1)[1]
                    
                    # Extract text from paragraphs and other text elements
                    text_elements = []
                    
                    # Look for common text elements
                    for elem in root.iter():
                        if elem.tag in ['p', 'h', 'span', 'text:p', 'text:h', 'text:span']:
                            if elem.text:
                                text_elements.append(elem.text)
                            # Also get text from nested elements
                            for child in elem:
                                if child.text:
                                    text_elements.append(child.text)
                                if child.tail:
                                    text_elements.append(child.tail)
                    
                    text_content = '\n'.join([t.strip() for t in text_elements if t and t.strip()])
                    
                except KeyError:
                    logger.error("content.xml not found in ODT file")
                    return "Could not extract content from ODT file - invalid format"
                except ET.ParseError as e:
                    logger.error(f"Error parsing ODT XML: {e}")
                    return "Could not parse ODT file content"
            
            if not text_content.strip():
                logger.warning("No text extracted from ODT file")
                return "No text content could be extracted from this ODT file."
            
            logger.info(f"✅ Extracted {len(text_content)} characters from ODT")
            return self._clean_text(text_content)
            
        except zipfile.BadZipFile:
            logger.error("Invalid ODT file - not a valid ZIP archive")
            raise ValueError("Invalid ODT file format")
        except Exception as e:
            logger.error(f"Error processing ODT: {str(e)}", exc_info=True)
            raise
    
    def _clean_markdown_syntax(self, text: str) -> str:
        """Remove common markdown syntax for plain text fallback."""
        # Remove headers
        text = re.sub(r'^#{1,6}\s+', '', text, flags=re.MULTILINE)
        
        # Remove bold/italic
        text = re.sub(r'\*{1,2}([^*]+)\*{1,2}', r'\1', text)
        text = re.sub(r'_{1,2}([^_]+)_{1,2}', r'\1', text)
        
        # Remove links
        text = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', text)
        
        # Remove code blocks
        text = re.sub(r'```[^`]*```', '', text, flags=re.DOTALL)
        text = re.sub(r'`([^`]+)`', r'\1', text)
        
        return text
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize extracted text."""
        if not text:
            return ""
        
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Remove excessive newlines (more than 2 consecutive)
        text = re.sub(r'\n{3,}', '\n\n', text)
        
        # Strip leading/trailing whitespace
        text = text.strip()
        
        return text
    
    def chunk_content(self, content: str, file_name: str, chunk_size: int = 1000, chunk_overlap: int = 200) -> List[Dict[str, Any]]:
        if not content or not content.strip():
            logger.warning(f"No content to chunk for file {file_name}")
            return []
        
        chunks = []
        content = content.strip()
        
        if len(content) <= chunk_size:
            chunks.append({
                'content': content,
                'metadata': {
                    'file_name': file_name,
                    'chunk_index': 0,
                    'total_chunks': 1,
                    'char_count': len(content),
                    'created_at': datetime.now().isoformat()
                }
            })
            logger.info(f"Single chunk created for small content (size: {len(content)})")
            return chunks
        
        start = 0
        chunk_index = 0
        max_iterations = len(content) // (chunk_size - chunk_overlap) + 10  # Increased buffer
        iteration = 0
        
        while start < len(content) and iteration < max_iterations:
            iteration += 1
            end = min(start + chunk_size, len(content))
            
            if end < len(content):
                search_start = max(start + chunk_size - chunk_overlap, start + chunk_size // 2)
                
                found_break = False
                for punct in ['. ', '.\n', '! ', '!\n', '? ', '?\n']:
                    punct_pos = content.rfind(punct, search_start, end)
                    if punct_pos != -1:
                        proposed_end = punct_pos + len(punct)
                        if proposed_end - chunk_overlap > start:  # Ensure advance
                            end = proposed_end
                            found_break = True
                        break
                
                if not found_break:
                    para_pos = content.rfind('\n\n', search_start, end)
                    if para_pos != -1:
                        proposed_end = para_pos + 2
                        if proposed_end - chunk_overlap > start:
                            end = proposed_end
                            found_break = True
                
                if not found_break:
                    space_pos = content.rfind(' ', search_start, end)
                    if space_pos != -1:
                        proposed_end = space_pos + 1
                        if proposed_end - chunk_overlap > start:
                            end = proposed_end
            
            chunk_content = content[start:end].strip()
            
            if chunk_content:
                chunks.append({
                    'content': chunk_content,
                    'metadata': {
                        'file_name': file_name,
                        'chunk_index': chunk_index,
                        'char_start': start,
                        'char_end': end,
                        'char_count': len(chunk_content),
                        'created_at': datetime.now().isoformat()
                    }
                })
                chunk_index += 1
            
            if end >= len(content):
                break  # Exit after adding the last chunk
            
            next_start = end - chunk_overlap
            if next_start <= start:
                logger.warning(f"Chunking progress stalled at start={start}, end={end}; forcing advance")
                next_start = start + 1  # Minimal advance to prevent loop
            start = next_start
        
        if iteration >= max_iterations:
            logger.error(f"Chunking exceeded max iterations for {file_name}; possible infinite loop prevented")
        
        total_chunks = len(chunks)
        if total_chunks == 0:
            logger.error(f"No chunks created for {file_name} despite having content (length: {len(content)})")
        
        for chunk in chunks:
            chunk['metadata']['total_chunks'] = total_chunks
        
        logger.info(f"✂️ Split content into {total_chunks} chunks (size: {chunk_size}, overlap: {chunk_overlap})")
        return chunks

    def _extract_text_with_ocr(self, file_path: str) -> str:
        """Extract text from image-based PDF using OCR with configuration and preprocessing."""
        if not OCR_AVAILABLE:
            raise ValueError("OCR dependencies not installed. Cannot perform OCR.")
        
        try:
            logger.info("🔍 Starting OCR extraction...")
            
            dpi = getattr(self.config, 'OCR_DPI', 300)
            languages = getattr(self.config, 'OCR_LANGUAGES', 'eng')
            
            logger.info(f"📄 Converting PDF to images (DPI: {dpi})...")
            images = convert_from_path(file_path, dpi=dpi)
            logger.info(f"📄 Converted PDF to {len(images)} images")
            
            text = ""
            successful_pages = 0
            
            for i, image in enumerate(images):
                logger.info(f"🔎 Processing page {i+1}/{len(images)} with OCR (lang: {languages})...")
                
                try:
                    img_array = np.array(image)
                    gray = cv2.cvtColor(img_array, cv2.COLOR_RGB2GRAY)
                    
                    # Otsu's binarization for adaptive thresholding
                    _, thresh = cv2.threshold(gray, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
                    
                    # Noise removal
                    denoised = cv2.medianBlur(thresh, 3)
                    
                    # Dilation to thicken handwriting lines
                    kernel = np.ones((2,2), np.uint8)  # Adjust kernel for thickness
                    dilated = cv2.dilate(denoised, kernel, iterations=1)
                    
                    processed_image = Image.fromarray(dilated)
                    
                    # Use PSM 3 for fully automatic layout
                    config = '--psm 3 --oem 1'
                    page_text = pytesseract.image_to_string(processed_image, lang=languages, config=config)
                    
                    if page_text.strip():
                        text += f"\n--- Page {i+1} (OCR) ---\n"
                        text += page_text + "\n\n"
                        successful_pages += 1
                        logger.info(f"✅ Page {i+1}: Extracted {len(page_text)} characters")
                    else:
                        logger.warning(f"⚠️ Page {i+1}: No text extracted")
                except Exception as page_error:
                    logger.error(f"❌ Error processing page {i+1}: {str(page_error)}")
                    continue
            
            if not text.strip():
                logger.warning("❌ OCR completed but no text was extracted from any page")
                return "OCR processing completed but no readable text was found in the PDF."
            
            logger.info(f"✅ OCR completed: Total {len(text)} characters extracted from {successful_pages}/{len(images)} pages")
            return text.strip()
            
        except Exception as e:
            logger.error(f"❌ OCR extraction failed: {str(e)}", exc_info=True)
            raise

